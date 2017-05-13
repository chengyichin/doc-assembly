from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.dml import MSO_THEME_COLOR

document = Document('termsheet.docx')

search = "BGL"

for paragraph in document.paragraphs:
    if search in paragraph.text:
        inline = paragraph.runs
        # Loop added to work with runs (strings with same style)
        skipNext = False
        for i in range(len(inline)):
            if skipNext:
                skipNext = False
            elif search in inline[i].text:
                searchIdx = inline[i].text.find(search)
                searchLen = len(search)                
                # no need to move the runs
                if searchLen == len(inline[i].text):
                    print("no move ", paragraph.text)
                    inline[i].font.size = Pt(25)
                    inline[i].font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
                # move one run - case search on left
                elif searchIdx == 0:
                    print("one move search on left", paragraph.text)
                    paragraph.add_run()
                    new_inline = paragraph.runs
                    for j in reversed(range(i+2,len(new_inline))):
                        new_inline[j].text = new_inline[j-1].text
                    
                    new_inline[i+1].text = new_inline[i].text[searchLen:]
                    new_inline[i].text = search
                    new_inline[i].font.size = Pt(25)
                    new_inline[i].font.color.theme_color = MSO_THEME_COLOR.ACCENT_1                   
                # move one run - case search on right
                elif searchIdx > 0 and (len(inline[i].text)-searchIdx-searchLen) <= 0:
                    print("one move search on right", paragraph.text)
                    paragraph.add_run()
                    new_inline = paragraph.runs
                    for j in reversed(range(i+2,len(new_inline))):
                        new_inline[j].text = new_inline[j-1].text
                    
                    new_inline[i+1].text = new_inline[i].text[searchIdx:]
                    new_inline[i+1].font.size = Pt(25)
                    new_inline[i+1].font.color.theme_color = MSO_THEME_COLOR.ACCENT_1   
                    new_inline[i].text = new_inline[i].text[:searchIdx]
                    skipNext = True

				# move two runs
                else:
                    print("two moves ", paragraph.text)
                    print(searchIdx)
                    paragraph.add_run()
                    paragraph.add_run()
                    new_inline = paragraph.runs
                    for j in reversed(range(i+3,len(new_inline))):
                        new_inline[j].text = new_inline[j-2].text
                    
                    new_inline[i+2].text = new_inline[i].text[(searchIdx+searchLen):]
                    new_inline[i+1].text = search
                    new_inline[i+1].font.size = Pt(25)
                    new_inline[i+1].font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
                    new_inline[i].text = new_inline[i].text[:searchIdx]
                    skipNext = True               

                print(paragraph.text)

document.save('termsheet_mod.docx')

